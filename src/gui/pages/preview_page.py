# -*- coding: utf-8 -*-
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QPushButton, 
    QFileDialog, QScrollArea, QLabel,
    QHBoxLayout, QFrame, QSplitter,
    QDialog, QApplication, QVBoxLayout, QMenu
)
from PyQt6.QtCore import (
    Qt, QThread, pyqtSignal, QUrl, 
    QSize, QPoint, QMimeData  # 从 QtCore 导入 QMimeData
)
from PyQt6.QtGui import (
    QPixmap, QImage, QIcon,
    QDrag, QKeySequence, QAction
)
from pathlib import Path
import tempfile
import os
import win32com.client
import pythoncom
import fitz  # PyMuPDF
from src.gui.components.loading_indicator import LoadingIndicator
from src.utils.temp_manager import TempManager
import threading
import queue
import time

class PreviewWorker(QThread):
    """异步预览工作线程"""
    progress = pyqtSignal(int)
    finished = pyqtSignal(dict)
    error = pyqtSignal(str)

    def __init__(self, original_doc_path, formatted_doc_path):
        super().__init__()
        self.original_doc_path = original_doc_path
        self.formatted_doc_path = formatted_doc_path
        self._is_running = True
        print("预览工作线程已创建")

    def run(self):
        try:
            print("开始生成预览...")
            page_images = {}
            
            # 渲染原始文档
            try:
                if not self._is_running:
                    return
                    
                # 检查文件是否存在
                if not os.path.exists(self.original_doc_path):
                    raise Exception(f"找不到原始文档: {self.original_doc_path}")
                
                from docx import Document
                original_doc = Document(self.original_doc_path)
                
                # 渲染前检查字体和其他资源
                try:
                    from PIL import Image, ImageDraw, ImageFont
                    # 测试字体加载
                    test_font = ImageFont.load_default()
                except Exception as e:
                    raise Exception(f"初始化渲染环境失败: {str(e)}")
                
                # 渲染文档
                page_images.update(self._render_document(original_doc, "original"))
                print("原始文档渲染完成")
                
            except Exception as e:
                print(f"渲染原始文档失败: {str(e)}")
                self.error.emit(str(e))
                return
            
            # 渲染格式化文档
            try:
                if not self._is_running:
                    return
                    
                # 检查文件是否存在
                if not os.path.exists(self.formatted_doc_path):
                    raise Exception(f"找不到格式化文档: {self.formatted_doc_path}")
                
                formatted_doc = Document(self.formatted_doc_path)
                page_images.update(self._render_document(formatted_doc, "formatted"))
                print("格式化文档渲染完成")
                
            except Exception as e:
                print(f"渲染格式化文档失败: {str(e)}")
                self.error.emit(str(e))
                return
            
            if self._is_running:
                print("发送渲染结果...")
                self.finished.emit(page_images)
            
        except Exception as e:
            print(f"预览生成失败: {str(e)}")
            self.error.emit(str(e))
        finally:
            self._is_running = False
            print("预览工作线程结束")

    def _render_document(self, doc, prefix):
        """渲染文档为图像"""
        try:
            from PIL import Image, ImageDraw, ImageFont
            import io
            
            images = {}
            page_height = 1200
            page_width = int(page_height * 0.7)
            
            # 设置字体和布局参数
            font_size = 14  # 减小字体大小
            line_height = int(font_size * 1.5)
            margin_top = 60
            margin_bottom = 60
            margin_left = 60
            margin_right = 60
            
            # 加载字体
            try:
                font = ImageFont.load_default()
                for font_name in ["simsun.ttc", "simhei.ttf", "msyh.ttc", "arial.ttf"]:
                    try:
                        test_font = ImageFont.truetype(font_name, font_size)
                        font = test_font
                        break
                    except:
                        continue
            except Exception as e:
                print(f"加载字体失败，使用默认字体: {str(e)}")
            
            current_page = Image.new('RGB', (page_width, page_height), 'white')
            draw = ImageDraw.Draw(current_page)
            y_position = margin_top
            page_num = 0
            
            # 添加页眉
            header_text = "原始文档" if prefix == "original" else "格式化预览"
            draw.text((margin_left, 20), header_text, font=font, fill='gray')
            
            # 处理段落
            total_paragraphs = len(doc.paragraphs)
            for i, para in enumerate(doc.paragraphs):
                if not self._is_running:
                    return images
                
                text = para.text.strip()
                if not text:
                    y_position += line_height // 2
                    continue
                
                # 计算文本布局
                text_width = page_width - margin_left - margin_right
                wrapped_text = self._wrap_text(text, font, text_width)
                para_height = len(wrapped_text) * line_height
                
                # 检查是否需要新页面
                if y_position + para_height > page_height - margin_bottom:
                    self._add_page_number(draw, page_num + 1, page_width, page_height, font)
                    images[f"{prefix}_{page_num}"] = self._convert_pil_to_qpixmap(current_page)
                    
                    # 创建新页面
                    current_page = Image.new('RGB', (page_width, page_height), 'white')
                    draw = ImageDraw.Draw(current_page)
                    y_position = margin_top
                    page_num += 1
                    
                    # 添加页眉
                    draw.text((margin_left, 20), header_text, font=font, fill='gray')
                
                # 绘制文本
                for line in wrapped_text:
                    draw.text((margin_left, y_position), line, font=font, fill='black')
                    y_position += line_height
                
                y_position += line_height // 2
                self.progress.emit(int((i + 1) * 100 / total_paragraphs))
            
            # 保存最后一页
            if y_position > margin_top:
                self._add_page_number(draw, page_num + 1, page_width, page_height, font)
                images[f"{prefix}_{page_num}"] = self._convert_pil_to_qpixmap(current_page)
            
            return images
            
        except Exception as e:
            print(f"渲染文档失败: {str(e)}")
            raise e

    def _add_page_number(self, draw, page_num, page_width, page_height, font):
        """添加页码"""
        try:
            page_number_text = f"- {page_num} -"
            w = draw.textlength(page_number_text, font=font)
            draw.text((page_width/2 - w/2, page_height - 40), 
                     page_number_text, font=font, fill='gray')
        except:
            # 如果页码添加失败，忽略错误继续执行
            pass

    def _wrap_text(self, text, font, max_width):
        """将文本按宽度换行"""
        lines = []
        current_line = []
        current_width = 0
        
        # 修改字体大小测量方法
        def get_text_width(text):
            try:
                # 新版本 PIL
                return font.getlength(text)
            except AttributeError:
                try:
                    # 旧版本 PIL
                    return font.getsize(text)[0]
                except:
                    # 如果都失败了，使用估算
                    return len(text) * font.size * 0.6
        
        for word in text.split():
            word_width = get_text_width(word + ' ')
            if current_width + word_width <= max_width:
                current_line.append(word)
                current_width += word_width
            else:
                if current_line:  # 确保当前行不为空
                    lines.append(' '.join(current_line))
                current_line = [word]
                current_width = get_text_width(word)
        
        if current_line:  # 添加最后一行
            lines.append(' '.join(current_line))
        
        return lines if lines else [text]  # 如果没有分行，返回原文本

    def _convert_pil_to_qpixmap(self, pil_image):
        """将PIL图像转换为QPixmap"""
        from PIL import ImageQt
        return QPixmap.fromImage(ImageQt.ImageQt(pil_image))

class PreviewPage(QWidget):
    def __init__(self, main_window):
        super().__init__()
        self.main_window = main_window
        self.temp_manager = TempManager()
        self.preview_worker = None
        self.last_format_hash = None
        self._needs_reload = True
        self.init_ui()
        
        # 添加快捷键
        self.save_action = QAction("保存文档", self)
        self.save_action.setShortcut("Ctrl+S")
        self.save_action.triggered.connect(self.save_document)
        self.addAction(self.save_action)
    
    def init_ui(self):
        """初始化用户界面"""
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        
        # 添加标题区域
        title_container = QFrame()
        title_container.setFixedHeight(50)
        title_container.setStyleSheet("""
            QFrame {
                background-color: #ffffff;
                border-bottom: 1px solid #e0e0e0;
            }
        """)
        title_layout = QHBoxLayout(title_container)
        title_layout.setContentsMargins(30, 0, 30, 0)
        
        original_label = QLabel("原始文档")
        formatted_label = QLabel("格式化预览")
        label_style = """
            QLabel {
                color: #333333;
                font-size: 15px;
                font-weight: bold;
                padding: 5px 15px;
                border-radius: 4px;
                background-color: #f8f9fa;
            }
        """
        original_label.setStyleSheet(label_style)
        formatted_label.setStyleSheet(label_style)
        title_layout.addWidget(original_label)
        title_layout.addStretch()
        title_layout.addWidget(formatted_label)
        
        layout.addWidget(title_container)
        
        # 创建中央部件
        central_widget = QWidget()
        central_layout = QVBoxLayout(central_widget)
        central_layout.setContentsMargins(20, 20, 20, 20)
        central_layout.setSpacing(0)
        
        # 创建外层容器
        container_frame = QFrame()
        container_frame.setStyleSheet("""
            QFrame {
                background-color: #ffffff;
                border: 1px solid #e0e0e0;
                border-radius: 12px;
            }
        """)
        container_layout = QVBoxLayout(container_frame)
        container_layout.setContentsMargins(1, 1, 1, 1)
        container_layout.setSpacing(0)
        
        # 建分割视图
        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.setChildrenCollapsible(False)
        splitter.setHandleWidth(1)
        splitter.setStyleSheet("""
            QSplitter::handle {
                background-color: #e0e0e0;
            }
        """)
        
        # 创建滚动区域和容器
        # 原始文档视图
        self.original_scroll = QScrollArea()
        self.original_container = QWidget()
        self.original_container.setObjectName("scrollContainer")
        self.original_layout = QVBoxLayout(self.original_container)
        self.original_layout.setAlignment(Qt.AlignmentFlag.AlignHCenter)
        self.original_layout.setContentsMargins(15, 15, 15, 15)
        self.original_scroll.setWidget(self.original_container)
        self.original_scroll.setWidgetResizable(True)
        
        # 格式化后的视图
        self.formatted_scroll = QScrollArea()
        self.formatted_container = QWidget()
        self.formatted_container.setObjectName("scrollContainer")
        self.formatted_layout = QVBoxLayout(self.formatted_container)
        self.formatted_layout.setAlignment(Qt.AlignmentFlag.AlignHCenter)
        self.formatted_layout.setContentsMargins(15, 15, 15, 15)
        self.formatted_scroll.setWidget(self.formatted_container)
        self.formatted_scroll.setWidgetResizable(True)
        
        # 设置滚动区域样式
        scroll_style = """
            QScrollArea {
                background-color: #f8f9fa;
                border: none;
                border-radius: 12px;
                margin: 10px;
            }
            QWidget#scrollContainer {
                background-color: #f8f9fa;
                border-radius: 12px;
            }
            QScrollBar:vertical {
                border: none;
                background: #f0f0f0;
                width: 8px;
                margin: 10px 2px;
            }
            QScrollBar::handle:vertical {
                background: #c1c1c1;
                min-height: 30px;
                border-radius: 4px;
            }
            QScrollBar::handle:vertical:hover {
                background: #a8a8a8;
            }
            QScrollBar::add-line:vertical,
            QScrollBar::sub-line:vertical {
                height: 0px;
            }
            QScrollBar::add-page:vertical,
            QScrollBar::sub-page:vertical {
                background: none;
            }
            QScrollBar:horizontal {
                height: 0px;
            }
        """
        
        self.original_scroll.setStyleSheet(scroll_style)
        self.formatted_scroll.setStyleSheet(scroll_style)
        
        # 添加到分割视图
        splitter.addWidget(self.original_scroll)
        splitter.addWidget(self.formatted_scroll)
        
        # 添加分割视图到容器
        container_layout.addWidget(splitter)
        
        # 添加容器到中央部件
        central_layout.addWidget(container_frame)
        layout.addWidget(central_widget)
        
        # 添加保存按钮到右上角
        save_button = QPushButton("保存文档", self)
        save_button.setStyleSheet("""
            QPushButton {
                background-color: #0078d4;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-size: 13px;
            }
            QPushButton:hover {
                background-color: #106ebe;
            }
            QPushButton:pressed {
                background-color: #005a9e;
            }
        """)
        save_button.clicked.connect(self.save_document)
        
        # 添加按钮到标题局
        title_layout.addWidget(save_button)
        
        # 添加快捷键
        self.save_action = QAction("保存文档", self)
        self.save_action.setShortcut("Ctrl+S")
        self.save_action.triggered.connect(self.save_document)
        self.addAction(self.save_action)
        
        # 添加右键菜单
        self.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.customContextMenuRequested.connect(self.show_context_menu)
        
        # 添加窗口大小变化事件处理
        def resizeEvent(self, event):
            super().resizeEvent(event)
            # 更新悬浮按钮位置
            self.save_fab.move(
                self.width() - self.save_fab.width() - 20,
                self.height() - self.save_fab.height() - 20
            )
        
        # 启用拖放
        self.setAcceptDrops(True)
        
        # 添加拖放提示区域
        self.drag_hint = QLabel("将预览拖放到文件夹以保存", self)
        self.drag_hint.setStyleSheet("""
            QLabel {
                color: #666666;
                background-color: #f8f9fa;
                padding: 8px 16px;
                border-radius: 20px;
                font-size: 13px;
            }
        """)
        self.drag_hint.hide()
    
    def update_preview(self):
        """更新预览内容"""
        if not self.main_window.document:
            return
        
        try:
            # 显示加载指示器
            self.show_loading_indicators()
            
            # 准备临时文件
            original_docx = self.temp_manager.get_temp_path("original.docx")
            formatted_docx = self.temp_manager.get_temp_path("formatted.docx")
            
            # 确保目录存在
            import os
            os.makedirs(os.path.dirname(original_docx), exist_ok=True)
            os.makedirs(os.path.dirname(formatted_docx), exist_ok=True)
            
            # 复制和格式化文档
            try:
                # 直接使用python-docx打开原始文档
                from docx import Document
                
                # 打开原始文档
                try:
                    # 先保存原始文档
                    original_doc = Document(self.main_window.document.path)
                    original_doc.save(original_docx)
                    print("原始文档保存成功")
                    
                    # 创建格式化文档的副本
                    formatted_doc = Document(self.main_window.document.path)
                    
                    # 应用格式
                    formatter = self.main_window.formatter
                    formatter.doc = formatted_doc
                    formatter.format()
                    
                    # 保存格式化文档
                    formatted_doc.save(formatted_docx)
                    print("格式化文档保存成功")
                    
                except Exception as e:
                    print(f"文档处理失败: {str(e)}")
                    self.main_window.show_message(f"文档处理失败: {str(e)}", error=True)
                    return
                    
            except Exception as e:
                print(f"准备预览文档失败: {str(e)}")
                self.main_window.show_message(f"准备预览文档失败: {str(e)}", error=True)
                return
            
            # 如果有正在运行的预览任务，先停止它
            if hasattr(self, 'preview_worker') and self.preview_worker and self.preview_worker.isRunning():
                try:
                    self.preview_worker.stop()
                    self.preview_worker.wait()
                except Exception as e:
                    print(f"停止旧预览任务失败: {str(e)}")
            
            try:
                # 创建并启动预览工作线程
                self.preview_worker = PreviewWorker(original_docx, formatted_docx)
                self.preview_worker.progress.connect(self.update_progress)
                self.preview_worker.finished.connect(self.show_preview_images)
                self.preview_worker.error.connect(self.handle_preview_error)
                self.preview_worker.start()
                
            except Exception as e:
                error_msg = f"启动预览任务失败: {str(e)}"
                print(error_msg)
                self.main_window.show_message(error_msg, error=True)
            
        except Exception as e:
            error_msg = f"预览失败: {str(e)}"
            print(error_msg)
            self.main_window.show_message(error_msg, error=True)
            self.clear_loading_indicators()
    
    def show_loading_indicators(self):
        """显示加载指示器"""
        self.clear_layout(self.original_layout)
        self.clear_layout(self.formatted_layout)
        
        # 创建预览区域占位并保存加载示器的引用
        self.loading_indicators = []
        
        for layout in [self.original_layout, self.formatted_layout]:
            placeholder = QFrame()
            placeholder.setStyleSheet("""
                QFrame {
                    background-color: white;
                    border: 1px solid #e0e0e0;
                    border-radius: 12px;
                    margin: 15px;
                    box-shadow: 0 2px 8px rgba(0, 0, 0, 0.08);
                }
            """)
            placeholder.setMinimumHeight(800)
            
            placeholder_layout = QVBoxLayout(placeholder)
            placeholder_layout.setContentsMargins(0, 0, 0, 0)
            
            loading = LoadingIndicator(self)
            self.loading_indicators.append(loading)
            placeholder_layout.addWidget(loading, 0, Qt.AlignmentFlag.AlignCenter)
            loading.start()
            
            layout.addWidget(placeholder)
    
    def update_progress(self, value):
        """更新进度"""
        # 只在控制台显示进度
        print(f"加载进度: {value}%")
        # 不在状态栏显示"正在加载"文字
        # self.main_window.statusBar.showMessage(f"正在加载预览... {value}%")
    
    def show_preview_images(self, page_images):
        """显示预览图像"""
        try:
            # 移除加载指示器
            for loading in self.loading_indicators:
                loading.stop()
                loading.deleteLater()
            self.loading_indicators.clear()
            
            # 清除现有内容
            self.clear_layout(self.original_layout)
            self.clear_layout(self.formatted_layout)
            
            # 修改页面显示宽度计算
            scroll_width = self.original_scroll.width()
            page_width = int(scroll_width * 0.92)  # 减小宽度比例，留出滚动条空间
            
            # 分别处理原始文档和格式化文档的页面
            original_pages = sorted([k for k in page_images.keys() if k.startswith('original_')])
            formatted_pages = sorted([k for k in page_images.keys() if k.startswith('formatted_')])
            
            # 显示原始文档页面
            for page_key in original_pages:
                page_num = int(page_key.split('_')[1]) + 1  # 提取页码加1
                pixmap = page_images[page_key]
                
                # 缩放图片以适应宽度
                scaled_pixmap = pixmap.scaledToWidth(
                    page_width, 
                    Qt.TransformationMode.SmoothTransformation
                )
                
                # 创建页面容器
                container = self.create_page_container(scaled_pixmap, page_num)
                self.original_layout.addWidget(container)
            
            # 显示格式化文档页面
            for page_key in formatted_pages:
                page_num = int(page_key.split('_')[1]) + 1  # 提取页码并加1
                pixmap = page_images[page_key]
                
                # 缩放图片以适应宽度
                scaled_pixmap = pixmap.scaledToWidth(
                    page_width, 
                    Qt.TransformationMode.SmoothTransformation
                )
                
                # 创建页面容器
                container = self.create_page_container(scaled_pixmap, page_num)
                self.formatted_layout.addWidget(container)
            
            # 添加底部空白
            original_spacer = QWidget()
            original_spacer.setMinimumHeight(20)
            self.original_layout.addWidget(original_spacer)
            
            formatted_spacer = QWidget()
            formatted_spacer.setMinimumHeight(20)
            self.formatted_layout.addWidget(formatted_spacer)
            
            self.main_window.statusBar.showMessage("预览加载完成", 3000)
            print("预览加载完成")
            
        except Exception as e:
            error_msg = f"显示预览失败: {str(e)}"
            print(error_msg)
            self.main_window.show_message(error_msg, error=True)
            import traceback
            traceback.print_exc()  # 添加详细错误信息
    
    def handle_preview_error(self, error_msg):
        """处理预览错误"""
        print(f"预览错误: {error_msg}")
        self.main_window.show_message(f"预览失败: {error_msg}", error=True)
        # 清除加载指示器
        self.clear_loading_indicators()
    
    def clear_loading_indicators(self):
        """清除加载指示器"""
        try:
            if hasattr(self, 'loading_indicators'):
                for loading in self.loading_indicators:
                    loading.stop()
                    loading.deleteLater()
                self.loading_indicators.clear()
        except Exception as e:
            print(f"清除加载指示器失败: {str(e)}")
    
    def convert_word_to_pdf(self, docx_path, pdf_path):
        """将Word文档转换为PDF"""
        word_app = None
        doc = None
        try:
            pythoncom.CoInitialize()
            
            print(f"开始转换文档: {docx_path}")
            
            # 创建Word应用实例
            word_app = win32com.client.DispatchEx("Word.Application")
            word_app.Visible = False
            word_app.DisplayAlerts = False
            
            print("Word应用创建成功")
            
            # 等待Word用就绪
            time.sleep(1)
            
            # 打开文档
            try:
                doc = word_app.Documents.Open(
                    docx_path,
                    ReadOnly=True,
                    Visible=False,
                    ConfirmConversions=False
                )
                print("文档打开成功")
            except Exception as e:
                raise Exception(f"打开文档失败: {str(e)}")
            
            # 等待文档加载完成
            time.sleep(1)
            
            try:
                # 保存为PDF
                pdf_path = str(Path(pdf_path).resolve())  # 确保使用完整路径
                doc.SaveAs2(
                    FileName=pdf_path,
                    FileFormat=17,  # wdFormatPDF = 17
                    AddToRecentFiles=False,
                    ReadOnlyRecommended=True
                )
                print(f"PDF保存成功: {pdf_path}")
            except Exception as e:
                raise Exception(f"保存PDF失败: {str(e)}")
            
        except Exception as e:
            raise Exception(f"转换PDF失败: {str(e)}")
            
        finally:
            try:
                # 关闭文档
                if doc:
                    try:
                        doc.Close(SaveChanges=False)
                        print("文档已关闭")
                    except:
                        pass
                
                # 退出Word应用
                if word_app:
                    try:
                        word_app.Quit()
                        print("Word应用已退出")
                    except:
                        pass
                
                # 释放COM对象
                if doc:
                    del doc
                if word_app:
                    del word_app
                
            except Exception as cleanup_error:
                print(f"清理资源时出错: {cleanup_error}")
                
            finally:
                pythoncom.CoUninitialize()
    
    def save_document(self):
        """保存格式化后的文档"""
        if not self.main_window.document:
            self.main_window.show_message("没有可保存的文档", error=True)
            return
        
        try:
            # 获取保存路径
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "保存文档",
                "",
                "Word文档 (*.docx)"
            )
            
            if not file_path:  # 用户取消了保存
                return
            
            # 确保文件扩展名正确
            if not file_path.lower().endswith('.docx'):
                file_path += '.docx'
            
            # 获取格式化后的临文档路径
            formatted_docx = self.temp_manager.get_temp_path("formatted.docx")
            
            if os.path.exists(formatted_docx):
                # 复制格式化后的文档到目标位置
                import shutil
                shutil.copy2(formatted_docx, file_path)
                
                self.main_window.show_message(f"文档已保存至: {file_path}")
            else:
                # 如果找不到格式化后的文档，尝试保存原始文档
                self.main_window.document.save(file_path)
                self.main_window.show_message(f"文档已保存至: {file_path}")
                
        except Exception as e:
            error_msg = f"保存文档失败: {str(e)}"
            print(error_msg)
            self.main_window.show_message(error_msg, error=True)
    
    def cleanup(self):
        """清理临时文"""
        try:
            for file in os.listdir(self.temp_dir):
                os.remove(os.path.join(self.temp_dir, file))
            os.rmdir(self.temp_dir)
        except Exception as e:
            print(f"清理临时文件失败: {str(e)}")
    
    def resizeEvent(self, event):
        """处理窗口大小变化事件"""
        super().resizeEvent(event)
        
        # 重新计算文档宽度
        window_width = self.width()
        available_width = window_width - 60  # 减小边距
        doc_width = available_width // 2
        
        # 更新滚动区域宽度
        self.original_scroll.setMinimumWidth(doc_width)
        self.formatted_scroll.setMinimumWidth(doc_width)
        
        # 如果有预览内容，重新加载以适应新的大小
        if hasattr(self, 'preview_worker') and self.preview_worker:
            self.update_preview()
    
    def clear_layout(self, layout):
        """清除布局中的所有部件"""
        while layout.count():
            item = layout.takeAt(0)
            widget = item.widget()
            if widget:
                widget.deleteLater()
    
    def _show_error_preview(self, message):
        """显示错误提示"""
        error_dialog = QDialog(self)
        error_dialog.setWindowTitle("错误提示")
        error_dialog.setStyleSheet("""
            QDialog {
                background-color: #f8f9fa;
                border: 1px solid #dee2e6;
                padding: 20px;
                border-radius: 5px;
            }
        """)
        
        error_label = QLabel(message)
        error_label.setStyleSheet("""
            QLabel {
                color: #dc3545;
                font-size: 14px;
                font-weight: bold;
            }
        """)
        
        layout = QVBoxLayout()
        error_dialog.setLayout(layout)
        layout.addWidget(error_label)
        layout.addStretch()
        
        error_dialog.exec()
    
    def _calculate_format_hash(self):
        """计算当前格式的哈希值"""
        if not hasattr(self.main_window, 'formatter') or not self.main_window.formatter:
            return None
            
        import hashlib
        import json
        
        try:
            # 获取格式设置
            format_spec = self.main_window.formatter.format_spec
            # 转换为JSON字符串并排序键值
            format_str = json.dumps(format_spec, sort_keys=True)
            # 计算哈希值
            return hashlib.md5(format_str.encode()).hexdigest()
        except Exception as e:
            print(f"计算格式哈希失败: {str(e)}")
            return None
    
    def _preview_content_exists(self):
        """检查预览内容是否已存在"""
        original_pdf = self.temp_manager.get_temp_path("original.pdf")
        formatted_pdf = self.temp_manager.get_temp_path("formatted.pdf")
        return os.path.exists(original_pdf) and os.path.exists(formatted_pdf)
    
    def create_page_container(self, pixmap, page_num):
        """创建页面容器"""
        page_container = QFrame()
        page_container.setStyleSheet("""
            QFrame {
                background-color: white;
                border: 1px solid #e0e0e0;
                border-radius: 8px;
                margin: 5px;
                padding: 0px;
                box-shadow: 0 2px 8px rgba(0, 0, 0, 0.08);
            }
            QFrame:hover {
                box-shadow: 0 4px 12px rgba(0, 0, 0, 0.12);
            }
        """)
        
        page_layout = QVBoxLayout(page_container)
        page_layout.setContentsMargins(10, 10, 10, 10)
        page_layout.setSpacing(8)
        
        # 创建图片标签
        page_label = QLabel()
        page_label.setPixmap(pixmap)
        page_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        page_label.setStyleSheet("""
            QLabel {
                background-color: white;
                padding: 0px;
            }
        """)
        page_layout.addWidget(page_label)
        
        # 添加页码
        page_number = QLabel(f"第 {str(page_num)} 页")
        page_number.setAlignment(Qt.AlignmentFlag.AlignCenter)
        page_number.setStyleSheet("""
            QLabel {
                color: #666666;
                font-size: 12px;
                padding: 4px 8px;
                background-color: #f8f9fa;
                border-radius: 4px;
                margin-top: 5px;
            }
        """)
        page_layout.addWidget(page_number)
        
        return page_container
    
    def show_context_menu(self, position):
        """显示右键菜单"""
        context_menu = QMenu(self)
        context_menu.setStyleSheet("""
            QMenu {
                background-color: #ffffff;
                border: 1px solid #e0e0e0;
                border-radius: 8px;
                padding: 5px;
            }
            QMenu::item {
                padding: 8px 20px;
                border-radius: 4px;
            }
            QMenu::item:selected {
                background-color: #f0f9ff;
                color: #0078d4;
            }
        """)
        
        # 添加保存选项
        save_action = context_menu.addAction("保存文档")
        save_action.setIcon(QIcon(":/icons/save.png"))
        save_action.triggered.connect(self.save_document)
        
        # 显示菜单
        context_menu.exec(self.mapToGlobal(position))
    
    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self.drag_start_position = event.pos()

    def mouseMoveEvent(self, event):
        if not (event.buttons() & Qt.MouseButton.LeftButton):
            return
            
        if (event.pos() - self.drag_start_position).manhattanLength() < QApplication.startDragDistance():
            return

        drag = QDrag(self)
        mimedata = QMimeData()
        
        # 创建临时文件
        temp_file = self.temp_manager.get_temp_path("temp.docx")
        self.main_window.document.save(temp_file)
        
        # 设置放数据
        mimedata.setUrls([QUrl.fromLocalFile(temp_file)])
        drag.setMimeData(mimedata)
        
        # 显示拖放提示
        self.drag_hint.show()
        
        # 执行拖放
        result = drag.exec(Qt.DropAction.CopyAction)
        
        # 隐藏提示
        self.drag_hint.hide()
    
    def show_preview(self, text, format_settings=None):
        # 只有在需要重新加载时才执行预览
        if self._needs_reload:
            # 执行预览逻辑
            # ... 现有的预览代码 ...
            self._needs_reload = False  # 预览完成后重置标志
        
    def force_reload(self):
        """强制设置需要重新加载"""
        self._needs_reload = True
    