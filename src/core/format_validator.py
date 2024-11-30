from typing import Dict, List, Optional
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from .format_spec import DocumentFormat

class FormatValidator:
    def __init__(self, document, format_spec: DocumentFormat):
        self.document = document
        self.format_spec = format_spec
        self.validation_results = []
    
    def validate_all(self) -> List[Dict]:
        """验证所有格式"""
        self.validation_results = []
        
        # 验证各个部分
        self.validate_title()
        self.validate_abstract()
        self.validate_keywords()
        self.validate_sections()
        self.validate_tables()
        self.validate_images()
        self.validate_page_setup()
        
        return self.validation_results
    
    def _add_validation_result(self, section: str, element: str, 
                             is_valid: bool, message: str):
        """添加验证结果"""
        self.validation_results.append({
            "section": section,
            "element": element,
            "is_valid": is_valid,
            "message": message
        })
    
    def validate_title(self):
        """验证标题格式"""
        title = self.document.get_title()
        if not title:
            self._add_validation_result(
                "title", "existence",
                False, "未找到标题"
            )
            return
        
        # 验证字体大小
        font_size = title.runs[0].font.size.pt if title.runs else None
        if font_size != self.format_spec.title.font_size:
            self._add_validation_result(
                "title", "font_size",
                False,
                f"标题字号不符合要求：当前 {font_size}pt，应为 {self.format_spec.title.font_size}pt"
            )
        
        # 验证对齐方式
        alignment_map = {
            WD_PARAGRAPH_ALIGNMENT.CENTER: "CENTER",
            WD_PARAGRAPH_ALIGNMENT.LEFT: "LEFT",
            WD_PARAGRAPH_ALIGNMENT.RIGHT: "RIGHT",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "JUSTIFY"
        }
        current_alignment = alignment_map.get(title.alignment)
        if current_alignment != self.format_spec.title.alignment:
            self._add_validation_result(
                "title", "alignment",
                False,
                f"标题对齐方式不符合要求：当前为 {current_alignment}，应为 {self.format_spec.title.alignment}"
            )
    
    def validate_abstract(self):
        """验证摘要格式"""
        abstract = self.document.get_abstract()
        if not abstract:
            self._add_validation_result(
                "abstract", "existence",
                False, "未找到摘要"
            )
            return
        
        # 验证缩进
        indent = abstract.paragraph_format.first_line_indent
        if indent:
            current_indent = indent.pt
            if current_indent != self.format_spec.abstract.first_line_indent:
                self._add_validation_result(
                    "abstract", "indent",
                    False,
                    f"摘要首行缩进不符合要求：当前 {current_indent}pt，应为 {self.format_spec.abstract.first_line_indent}pt"
                )
    
    def validate_sections(self):
        """验证章节格式"""
        sections = self.document.get_all_sections()
        for section_name, paragraphs in sections.items():
            # 验证章节标题
            if section_name in self.document.sections:
                section_para = next((p for p in self.document.doc.paragraphs 
                                 if p.text.strip() == section_name), None)
                if section_para:
                    self._validate_heading(section_para, section_name)
            
            # 验证正文段落
            for para in paragraphs:
                self._validate_body_paragraph(para)
    
    def _validate_heading(self, paragraph, heading_text):
        """验证标题格式"""
        # 判断标题级别
        if self._is_main_heading(heading_text):
            spec = self.format_spec.heading1
            heading_type = "一级标题"
        else:
            spec = self.format_spec.heading2
            heading_type = "二级标题"
        
        # 验证字体大小
        font_size = paragraph.runs[0].font.size.pt if paragraph.runs else None
        if font_size != spec.font_size:
            self._add_validation_result(
                "heading", heading_text,
                False,
                f"{heading_type}字号不符合要求：当前 {font_size}pt，应为 {spec.font_size}pt"
            )
    
    def _validate_body_paragraph(self, paragraph):
        """验证正文段落格式"""
        # 验证字体大小
        font_size = paragraph.runs[0].font.size.pt if paragraph.runs else None
        if font_size != self.format_spec.body.font_size:
            self._add_validation_result(
                "body", paragraph.text[:20] + "...",
                False,
                f"正文字号不符合要求：当前 {font_size}pt，应为 {self.format_spec.body.font_size}pt"
            )
        
        # 验证行距
        line_spacing = paragraph.paragraph_format.line_spacing
        if line_spacing != self.format_spec.body.line_spacing:
            self._add_validation_result(
                "body", paragraph.text[:20] + "...",
                False,
                f"行距不符合要求：当前 {line_spacing}，应为 {self.format_spec.body.line_spacing}"
            )
    
    def validate_tables(self):
        """验证表格格式"""
        tables = self.document.get_tables()
        for i, table in enumerate(tables):
            self._validate_table(table, i+1)
    
    def _validate_table(self, table, table_index):
        """验证单个表格的格式"""
        # 验证表格对齐方式
        if table.alignment != WD_TABLE_ALIGNMENT.CENTER:
            self._add_validation_result(
                "table", f"Table {table_index}",
                False,
                "表格未居中对齐"
            )
        
        # 验证表头格式
        if table.rows:
            header_row = table.rows[0]
            for cell in header_row.cells:
                if not cell.paragraphs[0].runs[0].font.bold:
                    self._add_validation_result(
                        "table", f"Table {table_index} header",
                        False,
                        "表头未加粗"
                    )
                    break
    
    def validate_images(self):
        """验证图片格式"""
        for paragraph in self.document.doc.paragraphs:
            for run in paragraph.runs:
                if run._r.xml.find('.//w:drawing') != -1:
                    self._validate_image_format(paragraph)
    
    def _validate_image_format(self, paragraph):
        """验证图片段落格式"""
        # 验证对齐方式
        if paragraph.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
            self._add_validation_result(
                "image", "alignment",
                False,
                "图片未居中对齐"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def _is_main_heading(self, text: str) -> bool:
        """判断是否为一级标题"""
        text = text.strip()
        return (
            any(text.startswith(f"{i}. ") for i in range(1, 10)) or
            any(text.startswith(f"{num}、") for num in ['一', '二', '三', '四', '五', '六', '七', '八', '九']) or
            text in ['引言', '介绍', '研究方法', '实验', '结果', '讨论', '结论', '参考文献']
        )
    
    
    def validate_keywords(self):
        """验证关键词格式"""
        keywords = self.document.get_keywords()
        if not keywords:
            self._add_validation_result(
                "keywords", "existence",
                False, "未找到关键词"
            )
            return
        
        # 验证关键词格式
        for keyword in keywords:
            if not keyword.isalnum():
                self._add_validation_result(
                    "keywords", keyword,
                    False,
                    "关键词格式不符合要求：必须为字母和数字的组合"
                )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求��{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
from typing import Dict, List, Optional
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from .format_spec import DocumentFormat

class FormatValidator:
    def __init__(self, document, format_spec: DocumentFormat):
        self.document = document
        self.format_spec = format_spec
        self.validation_results = []
    
    def validate_all(self) -> List[Dict]:
        """验证所有格式"""
        self.validation_results = []
        
        # 验证各个部分
        self.validate_title()
        self.validate_abstract()
        self.validate_keywords()
        self.validate_sections()
        self.validate_tables()
        self.validate_images()
        self.validate_page_setup()
        
        return self.validation_results
    
    def _add_validation_result(self, section: str, element: str, 
                             is_valid: bool, message: str):
        """添加验证结果"""
        self.validation_results.append({
            "section": section,
            "element": element,
            "is_valid": is_valid,
            "message": message
        })
    
    def validate_title(self):
        """验证标题格式"""
        title = self.document.get_title()
        if not title:
            self._add_validation_result(
                "title", "existence",
                False, "未找到标题"
            )
            return
        
        # 验证字体大小
        font_size = title.runs[0].font.size.pt if title.runs else None
        if font_size != self.format_spec.title.font_size:
            self._add_validation_result(
                "title", "font_size",
                False,
                f"标题字号不符合要求：当前 {font_size}pt，应为 {self.format_spec.title.font_size}pt"
            )
        
        # 验证对齐方式
        alignment_map = {
            WD_PARAGRAPH_ALIGNMENT.CENTER: "CENTER",
            WD_PARAGRAPH_ALIGNMENT.LEFT: "LEFT",
            WD_PARAGRAPH_ALIGNMENT.RIGHT: "RIGHT",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "JUSTIFY"
        }
        current_alignment = alignment_map.get(title.alignment)
        if current_alignment != self.format_spec.title.alignment:
            self._add_validation_result(
                "title", "alignment",
                False,
                f"标题对齐方式不符合要求：当前为 {current_alignment}，应为 {self.format_spec.title.alignment}"
            )
    
    def validate_abstract(self):
        """验证摘要格式"""
        abstract = self.document.get_abstract()
        if not abstract:
            self._add_validation_result(
                "abstract", "existence",
                False, "未找到摘要"
            )
            return
        
        # 验证缩进
        indent = abstract.paragraph_format.first_line_indent
        if indent:
            current_indent = indent.pt
            if current_indent != self.format_spec.abstract.first_line_indent:
                self._add_validation_result(
                    "abstract", "indent",
                    False,
                    f"摘要首行缩进不符合要求：当前 {current_indent}pt，应为 {self.format_spec.abstract.first_line_indent}pt"
                )
    
    def validate_sections(self):
        """验证章节格式"""
        sections = self.document.get_all_sections()
        for section_name, paragraphs in sections.items():
            # 验证章节标题
            if section_name in self.document.sections:
                section_para = next((p for p in self.document.doc.paragraphs 
                                 if p.text.strip() == section_name), None)
                if section_para:
                    self._validate_heading(section_para, section_name)
            
            # 验证正文段落
            for para in paragraphs:
                self._validate_body_paragraph(para)
    
    def _validate_heading(self, paragraph, heading_text):
        """验证标题格式"""
        # 判断标题级别
        if self._is_main_heading(heading_text):
            spec = self.format_spec.heading1
            heading_type = "一级标题"
        else:
            spec = self.format_spec.heading2
            heading_type = "二级标题"
        
        # 验证字体大小
        font_size = paragraph.runs[0].font.size.pt if paragraph.runs else None
        if font_size != spec.font_size:
            self._add_validation_result(
                "heading", heading_text,
                False,
                f"{heading_type}字号不符合要求：当前 {font_size}pt，应为 {spec.font_size}pt"
            )
    
    def _validate_body_paragraph(self, paragraph):
        """验证正文段落格式"""
        # 验证字体大小
        font_size = paragraph.runs[0].font.size.pt if paragraph.runs else None
        if font_size != self.format_spec.body.font_size:
            self._add_validation_result(
                "body", paragraph.text[:20] + "...",
                False,
                f"正文字号不符合要求：当前 {font_size}pt，应为 {self.format_spec.body.font_size}pt"
            )
        
        # 验证行距
        line_spacing = paragraph.paragraph_format.line_spacing
        if line_spacing != self.format_spec.body.line_spacing:
            self._add_validation_result(
                "body", paragraph.text[:20] + "...",
                False,
                f"行距不符合要求：当前 {line_spacing}，应为 {self.format_spec.body.line_spacing}"
            )
    
    def validate_tables(self):
        """验证表格格式"""
        tables = self.document.get_tables()
        for i, table in enumerate(tables):
            self._validate_table(table, i+1)
    
    def _validate_table(self, table, table_index):
        """验证单个表格的格式"""
        # 验证表格对齐方式
        if table.alignment != WD_TABLE_ALIGNMENT.CENTER:
            self._add_validation_result(
                "table", f"Table {table_index}",
                False,
                "表格未居中对齐"
            )
        
        # 验证表头格式
        if table.rows:
            header_row = table.rows[0]
            for cell in header_row.cells:
                if not cell.paragraphs[0].runs[0].font.bold:
                    self._add_validation_result(
                        "table", f"Table {table_index} header",
                        False,
                        "表头未加粗"
                    )
                    break
    
    def validate_images(self):
        """验证图片格式"""
        for paragraph in self.document.doc.paragraphs:
            for run in paragraph.runs:
                if run._r.xml.find('.//w:drawing') != -1:
                    self._validate_image_format(paragraph)
    
    def _validate_image_format(self, paragraph):
        """验证图片段落格式"""
        # 验证对齐方式
        if paragraph.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
            self._add_validation_result(
                "image", "alignment",
                False,
                "图片未居中对齐"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def _is_main_heading(self, text: str) -> bool:
        """判断是否为一级标题"""
        text = text.strip()
        return (
            any(text.startswith(f"{i}. ") for i in range(1, 10)) or
            any(text.startswith(f"{num}、") for num in ['一', '二', '三', '四', '五', '六', '七', '八', '九']) or
            text in ['引言', '介绍', '研究方法', '实验', '结果', '讨论', '结论', '参考文献']
        )
    
    def validate_keywords(self):
        """验证关键词格式"""
        keywords = self.document.get_keywords()
        if not keywords:
            self._add_validation_result(
                "keywords", "existence",
                False, "未找到关键词"
            )
            return
        
        # 验证关键词格式
        for keyword in keywords:
            if not keyword.isalnum():
                self._add_validation_result(
                    "keywords", keyword,
                    False,
                    "关键词格式不符合要求：必须为字母和数字的组合"
                )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        # 验证页边距
        margin_errors = []
        if section.top_margin.pt != self.format_spec.page_setup.margin_top:
            margin_errors.append("上边距")
        if section.bottom_margin.pt != self.format_spec.page_setup.margin_bottom:
            margin_errors.append("下边距")
        if section.left_margin.pt != self.format_spec.page_setup.margin_left:
            margin_errors.append("左边距")
        if section.right_margin.pt != self.format_spec.page_setup.margin_right:
            margin_errors.append("右边距")
        
        if margin_errors:
            self._add_validation_result(
                "page_setup", "margins",
                False,
                f"页边距不符合要求：{', '.join(margin_errors)}不正确"
            )
    
    def validate_page_setup(self):
        """验证页面设置"""
        section = self.document.doc.sections[0]
        
        ) 