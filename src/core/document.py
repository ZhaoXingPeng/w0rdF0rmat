from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from .ai_assistant import DocumentAI

class Document:
    def __init__(self, path):
        self.path = path
        self.doc = DocxDocument(path)
        self.ai_assistant = DocumentAI()
        # 存储论文各部分内容
        self.title = None
        self.abstract = None
        self.keywords = None
        self.sections = {}
        # 先尝试使用文档样式解析
        if not self._parse_by_styles():
            # 如果样式解析失败，尝试使用传统方法
            if not self._parse_document_traditional():
                # 如果传统方法也失败，最后才使用AI
                self._parse_with_ai()
    
    def _parse_by_styles(self) -> bool:
        """
        通过文档现有样式解析文档结构
        返回是否解析成功
        """
        try:
            # 检查文档是否有有效的样式结构
            styles = [s.name for s in self.doc.styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
            
            for para in self.doc.paragraphs:
                if not para.text.strip():
                    continue
                    
                style_name = para.style.name.lower()
                
                # 通过样式名称识别各个部分
                if 'title' in style_name or '标题' in style_name:
                    if not self.title:  # 只取第一个标题
                        self.title = para
                elif 'abstract' in style_name or '摘要' in style_name:
                    self.abstract = para
                elif 'keywords' in style_name or '关键词' in style_name:
                    self.keywords = para
                elif 'heading 1' in style_name or '标题 1' in style_name:
                    current_section = para.text
                    self.sections[current_section] = []
                elif current_section:
                    self.sections[current_section].append(para)
            
            # 如果至少识别出标题和一个章节，认为解析成功
            return bool(self.title and self.sections)
            
        except Exception as e:
            print(f"样式解析出错: {str(e)}")
            return False
    
    def _parse_document_traditional(self) -> bool:
        """
        使用传统方法解析文档结构
        返回是否解析成功
        """
        try:
            current_section = None
            found_structure = False
            
            for para in self.doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # 识别标题（通常是文档第一个非空段落）
                if not self.title:
                    self.title = para
                    found_structure = True
                    continue
                
                # 识别摘要部分
                if text.lower().startswith('abstract') or text.startswith('摘要'):
                    self.abstract = para
                    found_structure = True
                    continue
                
                # 识别关键词
                if text.lower().startswith('keywords') or text.startswith('关键词'):
                    self.keywords = para
                    found_structure = True
                    continue
                
                # 识别章节标题
                if self._is_section_heading(text):
                    current_section = text
                    self.sections[current_section] = []
                    found_structure = True
                elif current_section and text:
                    self.sections[current_section].append(para)
            
            return found_structure
            
        except Exception as e:
            print(f"传统解析方法出错: {str(e)}")
            return False
    
    def _parse_with_ai(self):
        """
        使用AI辅助解析文档结构
        """
        full_text = "\n".join([para.text for para in self.doc.paragraphs])
        ai_analysis = self.ai_assistant.analyze_document(full_text)
        
        if ai_analysis:
            self._update_structure_from_ai(ai_analysis)
    
    def _is_section_heading(self, text: str) -> bool:
        """
        判断是否为章节标题
        """
        # 检查数字编号格式（如：1. 2. 等）
        if any(text.startswith(f"{i}.") for i in range(1, 10)):
            return True
        
        # 检查中文数字编号格式（如：一、 二、 等）
        chinese_numbers = ['一', '二', '三', '四', '五', '六', '七', '八', '九']
        if any(text.startswith(f"{num}、") for num in chinese_numbers):
            return True
        
        # 检查特定的标题关键词
        heading_keywords = ['引言', '介绍', '研究方法', '实验', '结果', '讨论', '结论', '参考文献']
        return any(keyword in text for keyword in heading_keywords)
    
    def _update_structure_from_ai(self, ai_analysis):
        """
        根据AI分析结果更新文档结构
        """
        try:
            import json
            structure = json.loads(ai_analysis)
            
            # 更新文档各部分
            for para in self.doc.paragraphs:
                text = para.text.strip()
                
                # 根据AI识别结果匹配段落
                if text == structure.get('title'):
                    self.title = para
                elif text == structure.get('abstract'):
                    self.abstract = para
                elif text == structure.get('keywords'):
                    self.keywords = para
                
                # 处理章节
                for section in structure.get('sections', []):
                    if text == section['title']:
                        self.sections[text] = []
                        current_section = text
                    elif current_section:
                        self.sections[current_section].append(para)
                
        except Exception as e:
            print(f"解析AI结果时出错: {str(e)}")
            # 失败时使用传统方法解析
            self._parse_document_traditional()
    
    def get_title(self):
        """获取论文标题"""
        return self.title
    
    def get_abstract(self):
        """获取摘要部分"""
        return self.abstract
    
    def get_keywords(self):
        """获取关键词部分"""
        return self.keywords
    
    def get_section(self, section_name):
        """获取指定章节的内容"""
        return self.sections.get(section_name, [])
    
    def get_all_sections(self):
        """获取所有章节"""
        return self.sections
    
    def get_references(self):
        """获取参考文献部分"""
        for section_name, paragraphs in self.sections.items():
            if '参考文献' in section_name or 'references' in section_name.lower():
                return paragraphs
        return []
    
    def save(self, output_path):
        """保存文档"""
        self.doc.save(output_path)
    
    def get_paragraphs(self):
        """
        获取所有段落
        """
        return self.doc.paragraphs
    
    def get_tables(self):
        """
        获取所有表格
        """
        return self.doc.tables 
    
    def get_ai_format_suggestions(self, section_type):
        """
        获取AI对特定部分的格式建议
        """
        content = None
        if section_type == 'title':
            content = self.title.text if self.title else None
        elif section_type == 'abstract':
            content = self.abstract.text if self.abstract else None
        # ... 其他部分类似
        
        if content:
            return self.ai_assistant.suggest_formatting(section_type, content)
        return None